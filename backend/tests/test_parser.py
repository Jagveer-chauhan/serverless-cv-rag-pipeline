"""Tests for in-memory PyMuPDF PDF parsing and OCR fallback."""
import io
import pytest
try:
    import fitz
except ImportError:
    import pymupdf as fitz
from backend.app.services.parser import extract_text_from_pdf


def create_sample_pdf(text: str) -> bytes:
    """Helper to generate an in-memory PDF with specified text."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((50, 72), text, fontsize=11)
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes


def test_pymupdf_direct_extraction():
    sample_text = """
    Jane Doe
    Senior AI Architect | Python Specialist
    Email: jane.doe@example.com | Phone: +1 555-0199
    
    PROFESSIONAL SUMMARY
    Proven backend architect with 8+ years building high-throughput distributed systems and serverless RAG pipelines.
    
    WORK EXPERIENCE
    Lead Architect at CloudScale (2021 - Present)
    - Architected sub-5.0s CV parsing pipeline serving 1M+ monthly requests.
    - Optimized vector indexing latency using pgvector and asynchronous batching.
    
    EDUCATION
    M.S. in Computer Science - Stanford University
    B.S. in Software Engineering - MIT
    
    TECHNICAL SKILLS
    Python, FastAPI, Pydantic, SQLAlchemy, PostgreSQL, pgvector, Docker, PyMuPDF
    """
    pdf_bytes = create_sample_pdf(sample_text)
    extracted_text, meta = extract_text_from_pdf(pdf_bytes, filename="test_resume.pdf")

    assert "Jane Doe" in extracted_text
    assert "WORK EXPERIENCE" in extracted_text
    assert "CloudScale" in extracted_text
    assert meta["parser"] == "pymupdf_direct"
    assert meta["page_count"] == 1
    assert meta["char_count"] > 100
    assert meta["ocr_triggered"] is False


def test_empty_pdf_raises_error():
    with pytest.raises(ValueError, match="Document content.*is empty|PDF content is empty"):
        extract_text_from_pdf(b"", filename="empty.pdf")


def test_docx_extraction():
    from backend.app.services.parser import extract_text_from_docx, extract_text_from_document
    import docx

    doc = docx.Document()
    doc.add_heading("Alex Rivera - Senior Cloud Architect", 0)
    doc.add_paragraph("Email: alex.rivera@example.com | Skills: AWS, Python, Kubernetes")
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Company"
    table.rows[0].cells[1].text = "Role"
    table.rows[1].cells[0].text = "TechCorp"
    table.rows[1].cells[1].text = "Staff Engineer"

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Test direct docx extractor
    text, meta = extract_text_from_docx(docx_bytes, filename="alex_cv.docx")
    assert "Alex Rivera" in text
    assert "alex.rivera@example.com" in text
    assert "TechCorp" in text
    assert meta["parser"] == "python_docx"

    # Test unified document extractor
    text_uni, meta_uni = extract_text_from_document(docx_bytes, filename="alex_cv.docx")
    assert "Alex Rivera" in text_uni
    assert meta_uni["char_count"] > 30
