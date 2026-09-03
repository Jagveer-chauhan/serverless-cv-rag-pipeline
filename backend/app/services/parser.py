"""In-memory PDF & Word (DOCX/DOC) text extraction service using PyMuPDF and python-docx with OCR fallback."""
import io
import re
import zipfile
import xml.etree.ElementTree as ET
import logging
from typing import Tuple, Dict, Any, Optional

try:
    import pymupdf as fitz
except ImportError:
    import fitz

try:
    import docx
except ImportError:
    docx = None

from PIL import Image

logger = logging.getLogger("cv_rag_pipeline.parser")

MIN_TEXT_DENSITY_CHARS_PER_PAGE = 50


def extract_text_from_document(file_bytes: bytes, filename: str = "document.pdf") -> Tuple[str, Dict[str, Any]]:
    """Unified entry point to extract text from PDF, DOCX, or DOC bytes in-memory.

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    if not file_bytes:
        raise ValueError(f"Document content for '{filename}' is empty.")

    lower_name = filename.lower()
    if lower_name.endswith((".docx", ".doc")):
        return extract_text_from_docx(file_bytes, filename=filename)
    else:
        return extract_text_from_pdf(file_bytes, filename=filename)


def extract_text_from_docx(docx_bytes: bytes, filename: str = "document.docx") -> Tuple[str, Dict[str, Any]]:
    """Extracts text from DOCX (and DOC) bytes in-memory using python-docx with XML fallback."""
    if not docx_bytes:
        raise ValueError("DOCX content is empty.")

    # 1. Primary: Use python-docx
    if docx is not None:
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            sections_text = []

            # Extract paragraphs
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    sections_text.append(p_text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate adjacent duplicate cells from merged table cells
                    deduped_cells = []
                    for c in row_cells:
                        if not deduped_cells or deduped_cells[-1] != c:
                            deduped_cells.append(c)
                    if deduped_cells:
                        sections_text.append(" | ".join(deduped_cells))

            full_text = "\n\n".join(sections_text).strip()
            if full_text:
                metadata = {
                    "parser": "python_docx",
                    "page_count": 1,
                    "char_count": len(full_text),
                    "char_density": len(full_text),
                    "ocr_triggered": False,
                }
                logger.info(f"python-docx extracted {len(full_text)} chars from '{filename}'")
                return full_text, metadata
        except Exception as docx_err:
            logger.warning(f"python-docx parsing encountered error on '{filename}': {docx_err}. Trying zipfile XML fallback...")

    # 2. Secondary Fallback: Direct zipfile XML extraction (for .docx OpenXML)
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            if "word/document.xml" in zf.namelist():
                xml_content = zf.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                
                paragraphs = []
                for p in tree.iterfind(".//w:p", namespaces):
                    texts = [node.text for node in p.iterfind(".//w:t", namespaces) if node.text]
                    if texts:
                        paragraphs.append("".join(texts).strip())
                
                full_text = "\n\n".join(paragraphs).strip()
                if full_text:
                    metadata = {
                        "parser": "docx_zip_xml",
                        "page_count": 1,
                        "char_count": len(full_text),
                        "char_density": len(full_text),
                        "ocr_triggered": False,
                    }
                    logger.info(f"docx_zip_xml fallback extracted {len(full_text)} chars from '{filename}'")
                    return full_text, metadata
    except Exception as zip_err:
        logger.warning(f"Zipfile XML extraction failed on '{filename}': {zip_err}")

    # 3. Tertiary Fallback: Binary clean ASCII/UTF-8 string extraction (e.g. legacy .doc)
    try:
        # Extract printable ASCII/UTF-8 strings of length >= 3
        text_matches = re.findall(rb'[\x20-\x7E\r\n\t]{3,}', docx_bytes)
        extracted = "\n".join(m.decode("latin-1", errors="ignore").strip() for m in text_matches if m.strip())
        if extracted:
            metadata = {
                "parser": "binary_text_fallback",
                "page_count": 1,
                "char_count": len(extracted),
                "char_density": len(extracted),
                "ocr_triggered": False,
            }
            logger.info(f"binary_text_fallback extracted {len(extracted)} chars from '{filename}'")
            return extracted, metadata
    except Exception as bin_err:
        logger.error(f"Binary string fallback failed on '{filename}': {bin_err}")

    raise ValueError(f"Could not extract text from Word document '{filename}'.")


def extract_text_from_pdf(pdf_bytes: bytes, filename: str = "document.pdf") -> Tuple[str, Dict[str, Any]]:
    """Extracts text from PDF bytes in-memory using PyMuPDF with reading-order sorting and OCR fallback.
    
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    if not pdf_bytes:
        raise ValueError("PDF content is empty.")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_count = len(doc)
    if page_count == 0:
        raise ValueError("PDF document has 0 pages.")

    pages_text = []
    total_chars = 0

    for page_idx in range(page_count):
        page = doc.load_page(page_idx)
        # Sort blocks by geometric reading order (handles multi-column layouts)
        try:
            blocks = page.get_text("blocks", sort=True)
            # block[4] contains text in PyMuPDF block tuple (x0, y0, x1, y1, text, block_no, block_type)
            page_blocks_text = [b[4].strip() for b in blocks if len(b) > 4 and b[4].strip()]
            text = "\n\n".join(page_blocks_text) if page_blocks_text else (page.get_text("text", sort=True) or "")
        except Exception:
            text = page.get_text("text", sort=True) or ""

        pages_text.append(text)
        total_chars += len(text.strip())

    char_density = total_chars / page_count
    logger.info(
        f"PyMuPDF extracted {total_chars} chars across {page_count} pages "
        f"(density: {char_density:.1f} chars/page) for '{filename}'"
    )

    # Trigger OCR fallback if text density is below threshold
    if char_density < MIN_TEXT_DENSITY_CHARS_PER_PAGE or total_chars < 30:
        logger.warning(
            f"Low text density detected ({char_density:.1f} < {MIN_TEXT_DENSITY_CHARS_PER_PAGE}). "
            f"Triggering pytesseract OCR fallback for '{filename}'..."
        )
        ocr_text, ocr_meta = _extract_via_ocr_fallback(doc, filename)
        doc.close()
        return ocr_text, ocr_meta

    full_text = "\n\n".join(pages_text).strip()
    doc.close()

    metadata = {
        "parser": "pymupdf_direct",
        "page_count": page_count,
        "char_count": len(full_text),
        "char_density": round(char_density, 2),
        "ocr_triggered": False,
    }
    return full_text, metadata


def _extract_via_ocr_fallback(doc: fitz.Document, filename: str) -> Tuple[str, Dict[str, Any]]:
    """Fallback OCR extraction using pytesseract when direct text extraction fails."""
    try:
        import pytesseract
    except ImportError:
        logger.error("pytesseract is not installed. Returning partial direct text.")
        return "", {"parser": "fallback_unavailable", "ocr_triggered": True, "error": "pytesseract missing"}

    ocr_pages_text = []
    total_ocr_chars = 0
    page_count = len(doc)

    for page_idx in range(page_count):
        page = doc.load_page(page_idx)
        # Render page to high-res pixmap in memory (150 DPI for fast OCR)
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        try:
            page_text = pytesseract.image_to_string(img)
        except Exception as ocr_err:
            logger.warning(f"OCR failed on page {page_idx + 1} of '{filename}': {ocr_err}")
            page_text = ""
        ocr_pages_text.append(page_text)
        total_ocr_chars += len(page_text.strip())

    full_ocr_text = "\n\n".join(ocr_pages_text).strip()
    metadata = {
        "parser": "pytesseract_ocr",
        "page_count": page_count,
        "char_count": len(full_ocr_text),
        "char_density": round(total_ocr_chars / max(page_count, 1), 2),
        "ocr_triggered": True,
    }
    logger.info(f"pytesseract OCR completed for '{filename}' with {total_ocr_chars} characters.")
    return full_ocr_text, metadata
